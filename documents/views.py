# documents/views.py (corrected and enhanced: Document ni Order ga birlashtirish, QR va signed_at o'zgartirish)
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import JsonResponse, FileResponse
from django.db.models import Q
from django.views.decorators.http import require_http_methods
from django.views.decorators.csrf import csrf_exempt
from datetime import datetime
from django.utils import timezone
from .models import OrderSignature, Notification, Order, OrderSigner, AdditionalDocument, AdditionalDocumentTemplate
from users.models import CustomUser, Branch
from .forms import EditSignatureForm
import qrcode
from io import BytesIO
from django.core.files import File
import json
import os
import copy

# ReportLab imports for PDF generation
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Kirillcha shriftni ro'yxatdan o'tkazish (Windows va Linux/Docker)
_font_registered = False
for _font_path in [
    'C:/Windows/Fonts/arial.ttf',
    '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
    '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
]:
    try:
        import os as _os
        if _os.path.exists(_font_path):
            pdfmetrics.registerFont(TTFont('Arial', _font_path))
            _font_registered = True
            break
    except Exception:
        continue


@login_required
def branch_documents(request, branch_id):
    if request.user.role not in ['admin', 'director']:
        messages.error(request, "Ruxsat yo'q")
        return redirect('dashboard')

    branch = get_object_or_404(Branch, id=branch_id)

    selected_type = request.GET.get('type')

    orders = Order.objects.filter(branch=branch)

    # Tur bo‘yicha filter (I/CH, SH/T, buyruq)
    if selected_type in ['internal', 'external', 'official']:
        orders = orders.filter(document_type=selected_type)

    orders = orders.select_related(
        'branch',
        'created_by'
    ).prefetch_related(
        'signatures__user'
    )

    context = {
        'branch': branch,
        'documents': orders,
        'selected_type': selected_type,
    }

    return render(request, 'documents/branch_documents.html', context)


@login_required
def create_document(request):
    if request.user.role != 'admin':
        messages.error(request, 'Ruxsat yo\'q')
        return redirect('dashboard')

    if request.method == 'POST':
        try:
            title = request.POST.get('title')
            branch_id = request.POST.get('branch')
            document_type = request.POST.get('document_type', 'internal')
            signature_type = request.POST.get('signature_type')
            
            # Frontenddan kelgan ma'lumotni to'g'ri o'qish (array yoki vergulli string)
            employee_ids_raw = request.POST.getlist('employees') or request.POST.getlist('employees[]')
            employee_ids = []
            for item in employee_ids_raw:
                if ',' in item:
                    employee_ids.extend([x.strip() for x in item.split(',') if x.strip()])
                else:
                    if item.strip():
                        employee_ids.append(item.strip())
                        
            file = request.FILES.get('file')
            deadline_str = request.POST.get('deadline')
            number = request.POST.get('number')

            branch = get_object_or_404(Branch, id=branch_id)

            deadline = None
            if deadline_str:
                try:
                    deadline = datetime.strptime(deadline_str, '%Y-%m-%d').date()
                except ValueError:
                    messages.warning(request, "Muddati noto'g'ri formatda kiritildi, saqlanmadi")

            order = Order.objects.create(
                title=title,
                branch=branch,
                document_type=document_type,
                file=file,
                created_by=request.user,
                signature_type=signature_type,
                deadline=deadline,
                number=number
            )

            # Qo'shimcha hujjatlarni nomlari orqali yaratish (faylsiz)
            new_doc_names = request.POST.getlist('new_additional_names')
            new_doc_signers = request.POST.getlist('new_additional_signers')
            selected_saved_names = request.POST.getlist('selected_saved_doc_names')
            
            # Yangi kiritilgan nomlarni saqlash
            for idx, name in enumerate(new_doc_names):
                if name.strip():
                    try:
                        signer_id = new_doc_signers[idx] if idx < len(new_doc_signers) else None
                        signer = CustomUser.objects.filter(id=signer_id).first() if signer_id else None
                        new_doc = AdditionalDocument.objects.create(
                            name=name.strip(),
                            file=None,
                            branch=branch,
                            signer=signer
                        )
                        order.additional_docs.add(new_doc)
                    except Exception as e:
                        print(f"Error saving new additional document name: {e}")

            # Avvalgi saqlangan nomlarni saqlash
            for name in selected_saved_names:
                if name.strip():
                    try:
                        signer_id = request.POST.get(f'saved_doc_signers_{name}')
                        signer = CustomUser.objects.filter(id=signer_id).first() if signer_id else None
                        new_doc = AdditionalDocument.objects.create(
                            name=name.strip(),
                            file=None,
                            branch=branch,
                            signer=signer
                        )
                        order.additional_docs.add(new_doc)
                    except Exception as e:
                        print(f"Error saving selected additional document name: {e}")

            # Imzolar yaratish
            for i, emp_id in enumerate(employee_ids, start=1):
                try:
                    employee = CustomUser.objects.get(id=emp_id, role='employee')
                    OrderSigner.objects.create(
                        order=order,
                        user=employee,
                        order_number=i if signature_type == 'sequential' else 1,
                        required=True
                    )
                    OrderSignature.objects.create(
                        order=order,
                        user=employee,
                        order_number=i if signature_type == 'sequential' else 1,
                        signed=False
                    )

                    Notification.objects.create(
                        user=employee,
                        notification_type='signature_request',
                        order=order,
                        message=f"Sizdan \"{order.title}\" hujjatini imzolashingiz so'ralmoqda"
                    )
                except CustomUser.DoesNotExist:
                    continue

            messages.success(request, 'Hujjat muvaffaqiyatli yaratildi')
            return redirect('order_detail', order_id=order.id)

        except Exception as e:
            messages.error(request, f'Xatolik yuz berdi: {str(e)}')

    branches = Branch.objects.filter(parent_branch__isnull=True)
    employees = CustomUser.objects.values('id', 'first_name', 'last_name', 'middle_name', 'position', 'role')
    
    # Barcha mavjud unikal qo'shimcha hujjat shablonlarini olish
    saved_doc_names = AdditionalDocumentTemplate.objects.filter(is_active=True).values_list('name', flat=True).order_by('name')

    context = {
        'branches': branches,
        'employees': employees,
        'saved_doc_names': saved_doc_names,
    }
    return render(request, 'documents/create_document.html', context)

@login_required
def document_detail(request, order_id):
    order = get_object_or_404(
        Order.objects.select_related('branch', 'created_by')
                       .prefetch_related('signatures__user'),
        pk=order_id
    )

    can_view = (
        request.user.role == 'admin' or
        request.user.role == 'director' or
        request.user == order.created_by or
        order.signatures.filter(user=request.user).exists()
    )

    if not can_view:
        messages.error(request, 'Ushbu hujjatni ko\'rishga ruxsat yo\'q')
        return redirect('dashboard')

    signatures = order.signatures.all()
    signature_status = [
        {
            'employee': sig.user.get_full_name(),
            'position': sig.order_number,
            'signed': sig.signed,
            'signed_at': sig.signed_at,
            'comment': sig.comment,
            'has_signature_image': bool(sig.signature_image),
            'has_qr_code': bool(sig.qr_code)
        }
        for sig in signatures
    ]

    employees = CustomUser.objects.filter(role='employee')
    
    context = {
        'document': order,
        'signatures': signatures,
        'signature_status': signature_status,
        'employees': employees,
    }
    return render(request, 'documents/document_detail.html', context)

@login_required
def sign_document(request, signature_id):
    signature = get_object_or_404(OrderSignature, id=signature_id, user=request.user)

    if signature.signed:
        messages.warning(request, 'Siz bu hujjatni allaqachon imzolagansiz')
        return redirect('dashboard')

    # Ketma-ket imzolash tekshiruvi
    if signature.order.signature_type == 'sequential':
        pending_previous = OrderSignature.objects.filter(
            order=signature.order,
            order_number__lt=signature.order_number,
            signed=False
        ).exists()

        if pending_previous:
            messages.error(request, 'Oldingi imzolar hali qo‘yilmagan')
            return redirect('dashboard')

    if request.method == 'POST':
        if 'signature_image' in request.FILES:
            signature.signature_image = request.FILES['signature_image']

        signature.comment = request.POST.get('comment', '').strip()
        signature.signed = True
        signature.signed_at = datetime.now()
        signature.save()

        # Adminlarga xabar
        admins = CustomUser.objects.filter(role='admin')
        for admin in admins:
            Notification.objects.create(
                user=admin,
                notification_type='document_signed',
                order=signature.order,
                message=f"{request.user.get_full_name()} \"{signature.order.title}\" hujjatini imzoladi"
            )

        messages.success(request, 'Hujjat muvaffaqiyatli imzolandi!')
        return redirect('dashboard')

    return render(request, 'documents/sign_document.html', {'signature': signature})

@login_required
def get_document_types(request):
    branch_id = request.GET.get('branch_id')
    if not branch_id:
        return JsonResponse({'document_types': []})

    types_qs = DocumentType.objects.filter(branch_id=branch_id)
    data = [{'id': t.id, 'name': t.name} for t in types_qs]

    return JsonResponse({'document_types': data})

@login_required
def get_branch_employees(request):
    branch_id = request.GET.get('branch_id')
    if not branch_id:
        return JsonResponse({'employees': []})

    employees = CustomUser.objects.filter(
        branch__id=branch_id,
        role='employee'
    ).distinct().values('id', 'first_name', 'last_name', 'username', 'role')  # Rol qo'shildi

    return JsonResponse({'employees': list(employees)})

@login_required
def document_tracking(request, order_id):
    if request.user.role not in ['admin', 'director']:
        messages.error(request, 'Ruxsat yo\'q')
        return redirect('dashboard')

    order = get_object_or_404(
        Order.objects.select_related('branch', 'created_by')
                       .prefetch_related('signatures__user'),
        pk=order_id
    )

    signatures = order.signatures.order_by('order_number')

    signature_history = [
        {
            'employee': sig.user.get_full_name(),
            'position': sig.order_number,
            'status': 'Imzolandi' if sig.signed else 'Kutilmoqda',
            'signed_at': sig.signed_at.strftime('%Y-%m-%d %H:%M') if sig.signed_at else None,
            'comment': sig.comment or '—',
        }
        for sig in signatures
    ]

    context = {
        'document': order,
        'signature_history': signature_history,
    }
    return render(request, 'documents/document_tracking.html', context)


@login_required
def edit_order(request, order_id):
    if request.user.role != 'admin':
        messages.error(request, 'Ruxsat yo\'q')
        return redirect('dashboard')

    order = get_object_or_404(Order, id=order_id)

    if request.method == 'POST':
        title = request.POST.get('title')
        deadline_str = request.POST.get('deadline')

        order.title = title

        if deadline_str:
            try:
                order.deadline = datetime.strptime(deadline_str, '%Y-%m-%d').date()
            except ValueError:
                messages.warning(request, "Muddati noto'g'ri formatda kiritildi, saqlanmadi")
        else:
            order.deadline = None

        order.save()
        messages.success(request, 'Hujjat muvaffaqiyatli yangilandi')
        return redirect('document_detail', document_id=order.id)

    context = {
        'order': order,
    }
    return render(request, 'documents/edit_order.html', context)

@login_required
def edit_signature_time(request, signature_id):
    if request.user.role != 'admin':
        messages.error(request, 'Ruxsat yo\'q')
        return redirect('dashboard')
    
    signature = get_object_or_404(OrderSignature, id=signature_id)
    
    if request.method == 'POST':
        form = EditSignatureForm(request.POST, instance=signature)
        if form.is_valid():
            form.save()
            # QR kodni qayta generatsiya qilish
            signature.qr_code.delete(save=False)
            signature.save()  # save() ichida QR generatsiya qilinadi
            messages.success(request, 'Imzolangan vaqt o\'zgartirildi va QR kod yangilandi')
            return redirect('document_detail', document_id=signature.order.id)
    else:
        form = EditSignatureForm(instance=signature)
    
    context = {
        'form': form,
        'signature': signature,
    }
    return render(request, 'documents/edit_signature_time.html', context)


from qrcode.constants import ERROR_CORRECT_L
@login_required
@require_http_methods(["POST"])
def sign_with_fingerprint(request, signature_id):
    """Barmoq izi tasdiqlanganidan keyin hujjatni imzolash va QR kod yaratish"""
    signature = get_object_or_404(OrderSignature, id=signature_id, user=request.user)
    
    if signature.signed:
        return JsonResponse({'error': 'Bu hujjat allaqachon imzolangan'}, status=400)
    
    # Ketma-ket imzolash tekshiruvi
    if signature.order.signature_type == 'sequential':
        pending_previous = OrderSignature.objects.filter(
            order=signature.order,
            order_number__lt=signature.order_number,
            signed=False
        ).exists()
        
        if pending_previous:
            return JsonResponse({'error': 'Oldingi imzolar hali qo\'yilmagan'}, status=400)
    
    try:
        data = json.loads(request.body) if request.body else {}
    except json.JSONDecodeError:
        data = {}
    
    comment = data.get('comment', '')
    custom_time_str = data.get('custom_time')
    
    if custom_time_str:
        try:
            # Datetime local format is usually 'YYYY-MM-DDTHH:MM'
            signed_at = datetime.strptime(custom_time_str, '%Y-%m-%dT%H:%M')
        except ValueError:
            signed_at = datetime.now()
    else:
        signed_at = datetime.now()
    
    # QR kod ma'lumotlari
    user = request.user
    full_name = f"{user.last_name} {user.first_name} {user.middle_name}".strip()
    qr_data = (
        f"IMZO TASDIQLANGAN\n"
        f"F.I.O: {full_name}\n"
        f"Lavozim: {user.position or '-'}\n"
        f"Sarlavha: {signature.order.title}\n"
        f"Buyruq raqami: {signature.order.number}\n"
        f"Imzolangan: {signed_at.strftime('%d.%m.%Y %H:%M:%S')}\n"
        f"ID: {user.id}"
    )
    
    # QR kod generatsiya
    qr = qrcode.QRCode(
        version=None,
        error_correction=ERROR_CORRECT_L,
        box_size=4,
        border=2
    )

    qr.add_data(qr_data)
    qr.make(fit=True)

    qr_img = qr.make_image(fill_color="black", back_color="white")
    
    # QR kodni saqlash
    # qr_buffer = BytesIO()
    # qr_img.save(qr_buffer, format='PNG')
    # qr_buffer.seek(0)
    
    # Imzoni saqlash
    signature.signed = True
    signature.signed_at = signed_at
    signature.comment = comment
    
    # QR kodni ImageField ga saqlash
    qr_filename = f"qr_sign_{signature.order.number}_{user.username}_{signed_at.strftime('%Y%m%d%H%M%S')}.png"
    qr_save_buffer = BytesIO()
    qr_img.save(qr_save_buffer, format='PNG')
    qr_save_buffer.seek(0)
    signature.qr_code.save(qr_filename, File(qr_save_buffer), save=False)
    
    signature.save()
    
    # Word faylga alohida QR kod qo'shish bekor qilindi (PDF qatlamda qo'yiladi)
    word_embedded = False
    
    # Adminlarga xabar yuborish
    admins = CustomUser.objects.filter(role='admin')
    for admin in admins:
        Notification.objects.create(
            user=admin,
            notification_type='document_signed',
            order=signature.order,
            message=f"{user.get_full_name()} \"{signature.order.title}\" hujjatini barmoq izi bilan imzoladi"
        )
    
    # Javob
    response_data = {
        'success': True,
        'message': 'Hujjat muvaffaqiyatli imzolandi!',
        'word_updated': word_embedded,
    }
    
    if signature.qr_code:
        response_data['qr_code_url'] = signature.qr_code.url
    
    return JsonResponse(response_data)


@login_required
@require_http_methods(["POST"])
def sign_additional_doc_fingerprint(request, doc_id):
    """Barmoq izi orqali qo'shimcha hujjatni imzolash va QR kod yaratish"""
    doc = get_object_or_404(AdditionalDocument, id=doc_id, signer=request.user)
    
    if doc.is_signed:
        return JsonResponse({'error': 'Bu hujjat allaqachon imzolangan', 'success': False}, status=400)
    
    # QR kod ma'lumotlari
    user = request.user
    full_name = f"{user.last_name} {user.first_name} {user.middle_name}".strip()
    signed_at = datetime.now()
    
    qr_data = (
        f"QO'SHIMCHA HUJJAT IMZOLANGAN\n"
        f"F.I.O: {full_name}\n"
        f"Lavozim: {user.position or '-'}\n"
        f"Hujjat nomi: {doc.name}\n"
        f"Imzolangan: {signed_at.strftime('%d.%m.%Y %H:%M:%S')}\n"
        f"ID: {user.id}"
    )
    
    # QR kod generatsiya
    qr = qrcode.QRCode(
        version=None,
        error_correction=ERROR_CORRECT_L,
        box_size=4,
        border=2
    )

    qr.add_data(qr_data)
    qr.make(fit=True)

    qr_img = qr.make_image(fill_color="black", back_color="white")
    
    # Imzoni saqlash
    doc.is_signed = True
    doc.signed_at = signed_at
    
    # QR kodni ImageField ga saqlash
    qr_filename = f"qr_add_doc_{doc.id}_{user.username}_{signed_at.strftime('%Y%m%d%H%M%S')}.png"
    qr_save_buffer = BytesIO()
    qr_img.save(qr_save_buffer, format='PNG')
    qr_save_buffer.seek(0)
    doc.qr_code.save(qr_filename, File(qr_save_buffer), save=False)
    
    doc.save()
    
    # === Hujjatga xodim QR kodini (ostiga) bosish ===
    if doc.file:
        try:
            is_pdf = doc.file.name.lower().endswith('.pdf')
            is_word = doc.file.name.lower().endswith(('.doc', '.docx'))
            
            final_buffer = None
            if is_pdf:
                final_buffer = stamp_pdf_with_qrs(doc.file, doc.qr_code.path)
                stamped_filename = f"stamped_doc_{doc.id}_employee.pdf"
            elif is_word:
                final_buffer = stamp_word_with_qrs(doc.file, doc.qr_code.path)
                stamped_filename = f"stamped_doc_{doc.id}_employee.docx"
                
            if final_buffer:
                doc.stamped_file.save(stamped_filename, File(final_buffer), save=False)
                doc.save()
        except Exception as e:
            print(f"Error generating stamped additional doc (employee only) {doc.id}: {e}")
    
    # Javob
    response_data = {
        'success': True,
        'message': 'Qo\'shimcha hujjat muvaffaqiyatli imzolandi!',
    }
    
    return JsonResponse(response_data)

@login_required
def upload_stamped_pdf(request, order_id):
    """Pechat qo'yilgan PDF yuklash — yuklab olish QR kodi qo'shiladi."""
    import tempfile
    import os
    
    order = get_object_or_404(Order, id=order_id)
    
    if not order.director_approved:
        messages.error(request, "Faqat direktor tasdiqlagan hujjatlarga pechatli PDF yuklash mumkin.")
        return redirect('order_detail', order_id=order.id)
    
    if request.method != 'POST':
        return redirect('order_detail', order_id=order.id)
    
    uploaded_file = request.FILES.get('stamped_pdf')
    if not uploaded_file:
        messages.error(request, "PDF fayl tanlanmadi.")
        return redirect('order_detail', order_id=order.id)
    
    if not uploaded_file.name.lower().endswith('.pdf'):
        messages.error(request, "Faqat PDF format qabul qilinadi.")
        return redirect('order_detail', order_id=order.id)
    
    temp_files = []
    try:
        # Yuklangan PDFni vaqtinchalik saqlash
        fd, temp_pdf = tempfile.mkstemp(suffix='.pdf')
        os.close(fd)
        temp_files.append(temp_pdf)
        
        with open(temp_pdf, 'wb') as f:
            for chunk in uploaded_file.chunks():
                f.write(chunk)
        
        # Download QR kodi yaratish
        download_url = request.build_absolute_uri(f"/documents/download-stamped/{order.id}/")
        
        qr = qrcode.QRCode(version=1, box_size=5, border=1)
        qr.add_data(download_url)
        qr.make(fit=True)
        qr_img = qr.make_image(fill='black', back_color='white')
        
        fd, qr_path = tempfile.mkstemp(suffix='.png')
        os.close(fd)
        temp_files.append(qr_path)
        qr_img.save(qr_path, format='PNG')
        
        # QR kodni PDFga qo'shish (yuqori o'ng burchak)
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas as rl_canvas
        from PyPDF2 import PdfReader, PdfWriter
        
        # QR overlay yaratish
        fd, overlay_pdf = tempfile.mkstemp(suffix='.pdf')
        os.close(fd)
        temp_files.append(overlay_pdf)
        
        c = rl_canvas.Canvas(overlay_pdf, pagesize=A4)
        page_w, page_h = A4
        qr_size = 60  # pt
        margin = 30
        c.drawImage(qr_path, page_w - qr_size - margin, page_h - qr_size - margin, 
                     qr_size, qr_size)
        c.save()
        
        # Asl PDF + QR overlay birlashtirish
        reader = PdfReader(temp_pdf)
        overlay_reader = PdfReader(overlay_pdf)
        writer = PdfWriter()
        
        for i, page in enumerate(reader.pages):
            if i == 0:  # Faqat birinchi sahifaga QR qo'yish
                page.merge_page(overlay_reader.pages[0])
            writer.add_page(page)
        
        # Natijani saqlash
        fd, final_pdf = tempfile.mkstemp(suffix='.pdf')
        os.close(fd)
        temp_files.append(final_pdf)
        
        with open(final_pdf, 'wb') as f:
            writer.write(f)
        
        # Order ga saqlash
        from django.core.files import File
        with open(final_pdf, 'rb') as f:
            order.stamped_file.save(
                f"stamped_{order.number.replace('/', '_')}.pdf",
                File(f),
                save=True
            )
        
        messages.success(request, "Pechatli PDF muvaffaqiyatli yuklandi va yuklab olish QR kodi qo'shildi!")
        
    except Exception as e:
        print(f"upload_stamped_pdf error: {e}")
        import traceback
        traceback.print_exc()
        messages.error(request, f"Xatolik: {e}")
    finally:
        for tmp in temp_files:
            try:
                if os.path.exists(tmp):
                    os.remove(tmp)
            except Exception:
                pass
    
    return redirect('order_detail', order_id=order.id)


def download_stamped_pdf(request, order_id):
    """Pechatli PDF yuklab olish (login talab qilinmaydi — QR orqali ishlaydi)."""
    import os
    
    order = get_object_or_404(Order, pk=order_id)
    if not order.stamped_file:
        messages.error(request, "Pechatli fayl topilmadi")
        return redirect('order_detail', order_id=order.id)
    
    file_path = order.stamped_file.path
    if not os.path.exists(file_path):
        messages.error(request, "Fayl topilmadi.")
        return redirect('order_detail', order_id=order.id)
        
    with open(file_path, 'rb') as f:
        buffer = BytesIO(f.read())
        
    filename = f"pechatli_{order.number}.pdf".replace("/", "_")
    return FileResponse(buffer, as_attachment=True, filename=filename)

def download_additional_document(request, doc_id):
    """Qo'shimcha hujjatning elektron nusxasini (Pechatli PDF/Word) yuklab olish."""
    import os
    
    doc = get_object_or_404(AdditionalDocument, pk=doc_id)
    if not doc.stamped_file:
        messages.error(request, "Ushbu hujjat uchun pechatli fayl mavjud emas.")
        return redirect('dashboard')
        
    file_path = doc.stamped_file.path
    if not os.path.exists(file_path):
        messages.error(request, "Fayl topilmadi.")
        return redirect('dashboard')
        
    with open(file_path, 'rb') as f:
        buffer = BytesIO(f.read())
    
    ext = os.path.splitext(file_path)[1]
    filename = f"hujjat_elektron_{doc.id}{ext}"
    return FileResponse(buffer, as_attachment=True, filename=filename)


@login_required
def upload_additional_document_file(request, doc_id, order_id):
    if request.method == 'POST':
        doc = get_object_or_404(AdditionalDocument, id=doc_id)
        order = get_object_or_404(Order, id=order_id)
        
        # Security check
        can_edit = (
            request.user.role == 'admin' or
            request.user.role == 'director' or
            request.user == order.created_by
        )
        
        if not can_edit:
            messages.error(request, "Fayl yuklashga huquqingiz yo'q.")
            return redirect('order_detail', order_id=order.id)
            
        file = request.FILES.get('additional_file')
        if file:
            doc.file = file
            doc.save()
            messages.success(request, f"{doc.name} hujjatiga fayl muvaffaqiyatli yuklandi.")
        else:
            messages.error(request, "Iltimos, fayl tanlang.")
            
    return redirect('order_detail', order_id=order_id)


@login_required
def add_new_additional_document(request, order_id):
    if request.method == 'POST':
        order = get_object_or_404(Order, id=order_id)
        
        # Security check
        can_edit = (
            request.user.role == 'admin' or
            request.user.role == 'director' or
            request.user == order.created_by
        )
        
        if not can_edit:
            messages.error(request, "Qo'shimcha hujjat qo'shishga huquqingiz yo'q.")
            return redirect('order_detail', order_id=order.id)
            
        name = request.POST.get('name')
        file = request.FILES.get('file')
        signer_id = request.POST.get('signer')
        
        if name:
            try:
                signer = CustomUser.objects.filter(id=signer_id).first() if signer_id else None
                new_doc = AdditionalDocument.objects.create(
                    name=name,
                    file=file,
                    branch=order.branch,
                    signer=signer
                )
                order.additional_docs.add(new_doc)
                messages.success(request, "Yangi qo'shimcha hujjat muvaffaqiyatli qo'shildi.")
            except Exception as e:
                messages.error(request, f"Xatolik yuz berdi: {str(e)}")
        else:
            messages.error(request, "Hujjat nomi kiritilishi shart.")
            
    return redirect('order_detail', order_id=order_id)



def verify_document(request, order_id):
    """Hujjat imzolarini tekshirish — QR kodni skanerlash uchun (login kerak emas)."""
    order = get_object_or_404(Order, id=order_id)
    signatures = order.signatures.all().order_by('order_number')
    signed_count = signatures.filter(signed=True).count()
    total_count = signatures.count()
    
    return render(request, 'documents/verify_document.html', {
        'order': order,
        'signatures': signatures,
        'signed_count': signed_count,
        'total_count': total_count,
    })


def download_pdf(request, order_id):
    """Buyruqni PDF shaklida yuklab olish (login talab qilinmaydi — QR orqali ishlaydi).
    
    Jarayon:
    1. DOCX nusxasi yaratiladi
    2. QR kodlar Word faylga joylashtiriladi (yuklab olish QR yuqorida, imzo QR pastda)
    3. LibreOffice orqali PDF ga konvertatsiya qilinadi
    """
    import tempfile
    import os
    import shutil
    
    order = get_object_or_404(Order, id=order_id)
    
    if not order.file:
        messages.error(request, "Xatolik: Fayl topilmadi.")
        return redirect('order_detail', order_id=order.id)
    
    file_path = order.file.path
    if not os.path.exists(file_path):
        messages.error(request, "Xatolik: Fayl topilmadi.")
        return redirect('order_detail', order_id=order.id)
    
    temp_files = []
    
    try:
        if file_path.lower().endswith(('.doc', '.docx')):
            # STEP 1: Word nusxasiga QR kodlar joylash
            docx_with_qr = _embed_qr_in_docx(request, order, file_path, temp_files)
            
            # STEP 2: LibreOffice orqali PDF ga aylantirish
            pdf_path = _convert_docx_to_pdf(docx_with_qr, temp_files)
            
            with open(pdf_path, 'rb') as f:
                buffer = BytesIO(f.read())
            
            filename = f"hujjat_{order.number}.pdf".replace("/", "_")
            return FileResponse(buffer, as_attachment=True, filename=filename)
        
        elif file_path.lower().endswith('.pdf'):
            with open(file_path, 'rb') as f:
                buffer = BytesIO(f.read())
            filename = f"hujjat_{order.number}.pdf".replace("/", "_")
            return FileResponse(buffer, as_attachment=True, filename=filename)
        
        else:
            with open(file_path, 'rb') as f:
                buffer = BytesIO(f.read())
            ext = os.path.splitext(file_path)[1]
            filename = f"hujjat_{order.number}{ext}".replace("/", "_")
            return FileResponse(buffer, as_attachment=True, filename=filename)
        
    except Exception as e:
        print(f"download_pdf error: {e}")
        import traceback
        traceback.print_exc()
        messages.error(request, f"PDF yaratishda xatolik: {e}")
        return redirect('order_detail', order_id=order.id)
    finally:
        for tmp in temp_files:
            try:
                if os.path.isdir(tmp):
                    shutil.rmtree(tmp, ignore_errors=True)
                elif os.path.exists(tmp):
                    os.remove(tmp)
            except Exception:
                pass


@login_required
def download_docx(request, order_id):
    """DOCX faylni imzolangan holda yuklab olish — QR kodlar joylashtiriladi."""
    import os
    import tempfile
    import shutil
    
    order = get_object_or_404(Order, id=order_id)
    
    # Ruxsat tekshirish
    can_view = (
        request.user.role == 'admin' or
        request.user.role == 'director' or
        request.user == order.created_by or
        order.signatures.filter(user=request.user).exists()
    )
    
    if not can_view:
        messages.error(request, "Ushbu hujjatni yuklab olishga ruxsat yo'q")
        return redirect('dashboard')
    
    if not order.file:
        messages.error(request, "Xatolik: Fayl topilmadi.")
        return redirect('order_detail', order_id=order.id)
    
    file_path = order.file.path
    if not os.path.exists(file_path):
        messages.error(request, "Xatolik: Fayl topilmadi.")
        return redirect('order_detail', order_id=order.id)
    
    ext = os.path.splitext(file_path)[1].lower()
    filename = f"hujjat_{order.number}{ext}".replace("/", "_")
    
    # Agar DOCX bo'lmasa, oddiy yuklab olish
    if ext not in ['.docx']:
        with open(file_path, 'rb') as f:
            buffer = BytesIO(f.read())
        return FileResponse(buffer, as_attachment=True, filename=filename)
    
    # DOCX nusxasiga QR kodlar joylash (PDF bilan bir xil)
    temp_files = []
    try:
        temp_docx = _embed_qr_in_docx(request, order, file_path, temp_files)
        
        with open(temp_docx, 'rb') as f:
            buffer = BytesIO(f.read())
        
        return FileResponse(buffer, as_attachment=True, filename=filename)
        
    except Exception as e:
        print(f"download_docx error: {e}")
        import traceback
        traceback.print_exc()
        # Xato bo'lsa asl faylni yuborish
        with open(file_path, 'rb') as f:
            buffer = BytesIO(f.read())
        return FileResponse(buffer, as_attachment=True, filename=filename)
    finally:
        for tmp in temp_files:
            try:
                if os.path.exists(tmp):
                    os.remove(tmp)
            except Exception:
                pass

def _embed_qr_in_docx(request, order, docx_path, temp_files):
    """Word faylga QR kodlar joylash.
    
    - Yuqorida (header): Yuklab olish QR kodi
    - Pastda (oxirgi sahifa): Imzolovchilar QR kodlari
    
    Returns: QR joylashtirilgan DOCX faylning yo'li
    """
    import tempfile
    import os
    import shutil
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    
    # Nusxa yaratish
    fd, temp_docx = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    temp_files.append(temp_docx)
    shutil.copy2(docx_path, temp_docx)
    
    doc = Document(temp_docx)
    
    # === 1. HEADER: Yuklab olish QR kodi (Faqat 'application' turi uchun va direktor tasdiqlaganda) ===
    if order.document_type == 'application' and order.director_approved:
        try:
            download_url = request.build_absolute_uri(f"/documents/download-pdf/{order.id}/")
            qr = qrcode.QRCode(version=1, box_size=5, border=1)
            qr.add_data(download_url)
            qr.make(fit=True)
            img = qr.make_image(fill='black', back_color='white')
            
            fd, qr_path = tempfile.mkstemp(suffix='.png')
            os.close(fd)
            temp_files.append(qr_path)
            img.save(qr_path, format='PNG')
            
            # Headerga QR qo'yish
            section = doc.sections[0]
            header = section.header
            header.is_linked_to_previous = False
            
            # Header paragrafga QR rasm qo'shish
            if not header.paragraphs:
                hp = header.add_paragraph()
            else:
                hp = header.paragraphs[0]
            
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = hp.add_run()
            run.add_picture(qr_path, width=Inches(0.8))
        except Exception as e:
            print(f"Header QR error: {e}")
    
    # === 2. PASTDA: Imzolovchilar QR kodlari ===
    try:
        signatures = order.signatures.filter(signed=True).order_by('order_number')
        
        # Faqat direktor tasdiqlangandan keyin bitta imzo QR qo'yiladi
        if order.director_approved:
            # Imzolar tekshirish sahifasi URL
            verify_url = request.build_absolute_uri(f"/documents/verify/{order.id}/")
            
            doc.add_paragraph()
            
            # Bitta umumiy imzo QR
            sig_qr = qrcode.QRCode(version=None, box_size=6, border=1)
            sig_qr.add_data(verify_url)
            sig_qr.make(fit=True)
            sig_img = sig_qr.make_image(fill='black', back_color='white')
            
            fd, qr_img_path = tempfile.mkstemp(suffix='.png')
            os.close(fd)
            temp_files.append(qr_img_path)
            sig_img.save(qr_img_path, format='PNG')
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run()
            run.add_picture(qr_img_path, width=Inches(0.9))
            
            run2 = p.add_run("  Imzolangan va tasdiqlangan")
            run2.bold = True
            run2.font.size = Pt(9)
            run2.font.color.rgb = RGBColor(16, 124, 65)
            
            if order.director_approved_at:
                run3 = p.add_run(f"\n  {order.director_approved_at.strftime('%d.%m.%Y %H:%M')}")
                run3.font.size = Pt(8)
                run3.font.color.rgb = RGBColor(96, 94, 92)
    except Exception as e:
        print(f"Signature QR embedding error: {e}")
    
    doc.save(temp_docx)
    return temp_docx


def _convert_docx_to_pdf(docx_path, temp_files):
    """DOCX faylni PDF ga konvertatsiya qilish — faqat LibreOffice (Docker orqali kafolatlanadi)."""
    import tempfile
    import os
    import subprocess
    import shutil
    
    # LibreOffice yo'lini topish
    soffice_paths = [
        '/usr/bin/soffice',
        '/usr/bin/libreoffice',
        '/usr/lib/libreoffice/program/soffice',
        r'C:\Program Files\LibreOffice\program\soffice.exe',
        r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
    ]
    
    soffice = None
    for path in soffice_paths:
        if os.path.exists(path):
            soffice = path
            break
    
    if not soffice:
        soffice = shutil.which('soffice') or shutil.which('libreoffice')
    
    if not soffice:
        raise RuntimeError(
            "LibreOffice topilmadi! Dockerfile orqali o'rnatilganligini tekshiring."
        )
    
    print(f"LibreOffice: {soffice}")
    
    temp_dir = tempfile.mkdtemp()
    temp_files.append(temp_dir)
    
    docx_basename = os.path.basename(docx_path)
    temp_docx = os.path.join(temp_dir, docx_basename)
    shutil.copy2(docx_path, temp_docx)
    
    result = subprocess.run(
        [soffice, '--headless', '--norestore', '--convert-to', 'pdf',
         '--outdir', temp_dir, temp_docx],
        capture_output=True, text=True, timeout=120, cwd=temp_dir,
    )
    
    print(f"LibreOffice: returncode={result.returncode}, stdout={result.stdout}")
    if result.stderr:
        print(f"LibreOffice stderr: {result.stderr}")
    
    pdf_basename = os.path.splitext(docx_basename)[0] + '.pdf'
    output_pdf = os.path.join(temp_dir, pdf_basename)
    
    if os.path.exists(output_pdf) and os.path.getsize(output_pdf) > 0:
        fd, final_pdf = tempfile.mkstemp(suffix='.pdf')
        os.close(fd)
        temp_files.append(final_pdf)
        shutil.copy2(output_pdf, final_pdf)
        print(f"PDF yaratildi: {os.path.getsize(final_pdf)} bytes")
        return final_pdf
    
    raise RuntimeError(
        f"LibreOffice PDF yaratmadi. returncode={result.returncode}, "
        f"stderr={result.stderr}"
    )


def _emu_to_pt(emu):
    """EMU (English Metric Units) ni point (pt) ga aylantirish. 1 pt = 12700 EMU"""
    if emu is None:
        return None
    return emu / 12700.0


def _convert_docx_to_pdf_python(docx_path, temp_files):
    """DOCX ni PDF ga python-docx + reportlab orqali konvertatsiya qilish.
    
    Word hujjatdagi shrift, o'lcham, rang, alignment, rasmlar, marginlar
    va jadvallarni iloji boricha aniq ko'chiradi.
    """
    import tempfile
    import os
    from docx import Document
    from docx.shared import Pt, Emu
    from docx.oxml.ns import qn
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm, inch
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        Image as RLImage, KeepTogether, PageBreak
    )
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    
    # --- Shriftlarni ro'yxatdan o'tkazish ---
    _font_base = 'Helvetica'
    _font_bold = 'Helvetica-Bold'
    _font_italic = 'Helvetica-Oblique'
    _font_bold_italic = 'Helvetica-BoldOblique'
    
    # Arial shriftini topish va ro'yxatdan o'tkazish
    arial_paths = [
        'C:/Windows/Fonts/arial.ttf',
        'C:/Windows/Fonts/arialbd.ttf',
        'C:/Windows/Fonts/ariali.ttf',
        'C:/Windows/Fonts/arialbi.ttf',
        '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
        '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf',
        '/usr/share/fonts/truetype/liberation/LiberationSans-Italic.ttf',
        '/usr/share/fonts/truetype/liberation/LiberationSans-BoldItalic.ttf',
        '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
        '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
    ]
    
    try:
        # Windows: Arial
        if os.path.exists('C:/Windows/Fonts/arial.ttf'):
            pdfmetrics.registerFont(TTFont('DocFont', 'C:/Windows/Fonts/arial.ttf'))
            pdfmetrics.registerFont(TTFont('DocFont-Bold', 'C:/Windows/Fonts/arialbd.ttf'))
            pdfmetrics.registerFont(TTFont('DocFont-Italic', 'C:/Windows/Fonts/ariali.ttf'))
            pdfmetrics.registerFont(TTFont('DocFont-BoldItalic', 'C:/Windows/Fonts/arialbi.ttf'))
            _font_base = 'DocFont'
            _font_bold = 'DocFont-Bold'
            _font_italic = 'DocFont-Italic'
            _font_bold_italic = 'DocFont-BoldItalic'
        # Linux: Liberation Sans (metrically compatible with Arial)
        elif os.path.exists('/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf'):
            pdfmetrics.registerFont(TTFont('DocFont', '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf'))
            pdfmetrics.registerFont(TTFont('DocFont-Bold', '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf'))
            pdfmetrics.registerFont(TTFont('DocFont-Italic', '/usr/share/fonts/truetype/liberation/LiberationSans-Italic.ttf'))
            pdfmetrics.registerFont(TTFont('DocFont-BoldItalic', '/usr/share/fonts/truetype/liberation/LiberationSans-BoldItalic.ttf'))
            _font_base = 'DocFont'
            _font_bold = 'DocFont-Bold'
            _font_italic = 'DocFont-Italic'
            _font_bold_italic = 'DocFont-BoldItalic'
        # Linux: DejaVu Sans
        elif os.path.exists('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'):
            pdfmetrics.registerFont(TTFont('DocFont', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'))
            pdfmetrics.registerFont(TTFont('DocFont-Bold', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'))
            _font_base = 'DocFont'
            _font_bold = 'DocFont-Bold'
    except Exception as e:
        print(f"Shrift ro'yxatdan o'tkazishda xatolik: {e}")
    
    # PDF fayl yaratish
    fd, pdf_path = tempfile.mkstemp(suffix='.pdf')
    os.close(fd)
    temp_files.append(pdf_path)
    
    doc = Document(docx_path)
    
    # --- Word hujjatdan marginlarni olish ---
    left_margin = 25 * mm
    right_margin = 25 * mm
    top_margin = 20 * mm
    bottom_margin = 20 * mm
    
    try:
        section = doc.sections[0]
        if section.left_margin is not None:
            left_margin = _emu_to_pt(section.left_margin) * (mm / 2.835)  # pt -> mm -> reportlab
            left_margin = _emu_to_pt(section.left_margin)  # EMU -> pt, reportlab pt bilan ishlaydi
        if section.right_margin is not None:
            right_margin = _emu_to_pt(section.right_margin)
        if section.top_margin is not None:
            top_margin = _emu_to_pt(section.top_margin)
        if section.bottom_margin is not None:
            bottom_margin = _emu_to_pt(section.bottom_margin)
    except Exception as e:
        print(f"Margin o'qishda xatolik: {e}")
    
    # PDF hujjat sozlamalari — Word marginlari bilan
    pdf_doc = SimpleDocTemplate(
        pdf_path, pagesize=A4,
        leftMargin=left_margin,
        rightMargin=right_margin,
        topMargin=top_margin,
        bottomMargin=bottom_margin
    )
    
    available_width = A4[0] - left_margin - right_margin
    
    styles = getSampleStyleSheet()
    
    # --- Rasmlarni DOCX dan ajratib olish ---
    def _extract_images_from_paragraph(para, para_index):
        """Paragrafdan inline rasmlarni ajratib olish"""
        images = []
        try:
            # Inline shapes (rasmlar)
            for run in para.runs:
                run_xml = run._element
                drawing_elements = run_xml.findall(qn('w:drawing'))
                for drawing in drawing_elements:
                    # Inline rasm
                    inline = drawing.find(qn('wp:inline'))
                    if inline is None:
                        inline = drawing.find(qn('wp:anchor'))
                    if inline is None:
                        continue
                    
                    # Rasm o'lchamlari (EMU)
                    extent = inline.find(qn('wp:extent'))
                    img_width = None
                    img_height = None
                    if extent is not None:
                        cx = extent.get('cx')
                        cy = extent.get('cy')
                        if cx:
                            img_width = int(cx) / 914400.0 * inch  # EMU -> inch -> pt
                        if cy:
                            img_height = int(cy) / 914400.0 * inch
                    
                    # Rasm ma'lumotlarini olish
                    blip = inline.find('.//' + qn('a:blip'))
                    if blip is not None:
                        rId = blip.get(qn('r:embed'))
                        if rId:
                            try:
                                image_part = para.part.related_parts[rId]
                                image_data = image_part.blob
                                
                                # Vaqtinchalik fayl
                                img_ext = os.path.splitext(image_part.partname)[1] or '.png'
                                fd_img, img_path = tempfile.mkstemp(suffix=img_ext)
                                os.close(fd_img)
                                temp_files.append(img_path)
                                
                                with open(img_path, 'wb') as f:
                                    f.write(image_data)
                                
                                images.append({
                                    'path': img_path,
                                    'width': img_width,
                                    'height': img_height,
                                })
                            except Exception as e:
                                print(f"Rasm ajratishda xatolik: {e}")
        except Exception as e:
            print(f"Paragrafdan rasm olishda xatolik: {e}")
        return images
    
    def _get_paragraph_alignment(para):
        """Paragrafning alignment-ini aniqlash"""
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                return TA_CENTER
            elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                return TA_RIGHT
            elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                return TA_JUSTIFY
        except Exception:
            pass
        return TA_LEFT
    
    def _get_font_size_from_para(para):
        """Paragrafdan shrift o'lchamini aniqlash"""
        for run in para.runs:
            if run.font.size is not None:
                return _emu_to_pt(run.font.size)
        # Default
        return 12
    
    def _get_run_font_color(run):
        """Run-ning rangini olish"""
        try:
            if run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                return f'#{rgb}'
        except Exception:
            pass
        return None
    
    def _build_rich_text(para):
        """Paragrafdan rich text (HTML teg bilan) yaratish"""
        formatted_parts = []
        for run in para.runs:
            run_text = run.text
            if not run_text:
                continue
            
            # XML escape
            run_text = run_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # Shrift o'lchami
            font_size = None
            if run.font.size is not None:
                font_size = int(_emu_to_pt(run.font.size))
            
            # Rang
            color = _get_run_font_color(run)
            
            # Font nomi
            font_name = None
            if run.bold and run.italic:
                font_name = _font_bold_italic
            elif run.bold:
                font_name = _font_bold
            elif run.italic:
                font_name = _font_italic
            else:
                font_name = _font_base
            
            # ReportLab font tag
            parts = []
            if font_size or color or font_name:
                tag_attrs = []
                if font_name:
                    tag_attrs.append(f'face="{font_name}"')
                if font_size:
                    tag_attrs.append(f'size="{font_size}"')
                if color:
                    tag_attrs.append(f'color="{color}"')
                tag_str = ' '.join(tag_attrs)
                run_text = f'<font {tag_str}>{run_text}</font>'
            
            if run.bold:
                run_text = f'<b>{run_text}</b>'
            if run.italic:
                run_text = f'<i>{run_text}</i>'
            if run.underline:
                run_text = f'<u>{run_text}</u>'
            
            formatted_parts.append(run_text)
        
        return ''.join(formatted_parts) if formatted_parts else None
    
    story = []
    
    # --- Hujjat elementlarini ketma-ketlikda qayta ishlash ---
    # Word da paragraflar va jadvallar aralash keladi,
    # ularni to'g'ri tartibda qayta ishlash kerak
    
    body = doc.element.body
    
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        
        # === PARAGRAF ===
        if tag == 'p':
            from docx.text.paragraph import Paragraph as DocxParagraph
            para = DocxParagraph(child, doc)
            
            text = para.text.strip()
            
            # Sahifa uzilishi (page break) tekshirish
            has_page_break = False
            for run in para.runs:
                run_xml = run._element
                br_elements = run_xml.findall(qn('w:br'))
                for br in br_elements:
                    if br.get(qn('w:type')) == 'page':
                        has_page_break = True
                        break
            
            if has_page_break:
                story.append(PageBreak())
            
            # Rasmlarni tekshirish
            images = _extract_images_from_paragraph(para, 0)
            
            if not text and not images:
                # Bo'sh paragrafdan spacing
                space_after = 6
                try:
                    pf = para.paragraph_format
                    if pf.space_after is not None:
                        space_after = _emu_to_pt(pf.space_after)
                except Exception:
                    pass
                story.append(Spacer(1, max(space_after, 2)))
                continue
            
            # Paragraf formatting
            alignment = _get_paragraph_alignment(para)
            font_size = _get_font_size_from_para(para)
            leading = font_size * 1.35  # line spacing
            
            # Paragraph spacing
            space_before = 0
            space_after = 3
            left_indent = 0
            first_line_indent = 0
            
            try:
                pf = para.paragraph_format
                if pf.space_before is not None:
                    space_before = _emu_to_pt(pf.space_before)
                if pf.space_after is not None:
                    space_after = _emu_to_pt(pf.space_after)
                if pf.left_indent is not None:
                    left_indent = _emu_to_pt(pf.left_indent)
                if pf.first_line_indent is not None:
                    first_line_indent = _emu_to_pt(pf.first_line_indent)
                # Line spacing
                if pf.line_spacing is not None:
                    try:
                        if pf.line_spacing_rule is not None:
                            from docx.enum.text import WD_LINE_SPACING
                            if pf.line_spacing_rule == WD_LINE_SPACING.EXACTLY:
                                leading = _emu_to_pt(pf.line_spacing)
                            elif pf.line_spacing_rule == WD_LINE_SPACING.AT_LEAST:
                                leading = max(leading, _emu_to_pt(pf.line_spacing))
                            else:
                                # Multiple (proportional)
                                if isinstance(pf.line_spacing, float):
                                    leading = font_size * pf.line_spacing
                    except Exception:
                        pass
            except Exception:
                pass
            
            # Stil aniqlash
            style_name = (para.style.name or '').lower()
            is_heading = 'heading' in style_name or 'title' in style_name
            
            if is_heading:
                font_name = _font_bold
                if font_size <= 12:
                    font_size = max(font_size, 14)
            else:
                is_bold = all(run.bold for run in para.runs if run.text.strip()) if para.runs else False
                font_name = _font_bold if is_bold else _font_base
            
            # Dynamic ParagraphStyle
            style_key = f'dyn_{id(para)}'
            try:
                dyn_style = ParagraphStyle(
                    name=style_key,
                    fontName=font_name,
                    fontSize=font_size,
                    leading=leading,
                    spaceBefore=space_before,
                    spaceAfter=space_after,
                    alignment=alignment,
                    leftIndent=max(left_indent, 0),
                    firstLineIndent=first_line_indent,
                )
            except Exception:
                dyn_style = ParagraphStyle(
                    name=style_key,
                    fontName=_font_base,
                    fontSize=12,
                    leading=16,
                    spaceAfter=6,
                    alignment=TA_LEFT,
                )
            
            # Rich matn
            rich_text = _build_rich_text(para)
            if not rich_text:
                if text:
                    rich_text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                else:
                    rich_text = ''
            
            if rich_text:
                try:
                    story.append(Paragraph(rich_text, dyn_style))
                except Exception:
                    safe_text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(safe_text, dyn_style))
            
            # Rasmlarni qo'shish
            for img_info in images:
                try:
                    img_w = img_info.get('width')
                    img_h = img_info.get('height')
                    
                    # O'lchamni cheklash (sahifa kengligidan oshmasligi kerak)
                    if img_w and img_w > available_width:
                        ratio = available_width / img_w
                        img_w = available_width
                        if img_h:
                            img_h = img_h * ratio
                    
                    rl_img = RLImage(img_info['path'])
                    if img_w:
                        rl_img.drawWidth = img_w
                    if img_h:
                        rl_img.drawHeight = img_h
                    
                    # Rasmni sahifa kengligiga sig'dirish
                    if rl_img.drawWidth > available_width:
                        ratio = available_width / rl_img.drawWidth
                        rl_img.drawWidth = available_width
                        rl_img.drawHeight = rl_img.drawHeight * ratio
                    
                    story.append(rl_img)
                except Exception as e:
                    print(f"Rasm PDF ga qo'shishda xatolik: {e}")
        
        # === JADVAL ===
        elif tag == 'tbl':
            from docx.table import Table as DocxTable
            try:
                tbl = DocxTable(child, doc)
                table_data = []
                
                for row_idx, row in enumerate(tbl.rows):
                    row_data = []
                    for cell in row.cells:
                        # Hujayra ichidagi matnni formatlash
                        cell_parts = []
                        for cp in cell.paragraphs:
                            cp_text = cp.text.strip()
                            if not cp_text:
                                continue
                            
                            # Hujayra ichidagi run formatlarni saqlash
                            cell_rich = _build_rich_text(cp)
                            if cell_rich:
                                cell_parts.append(cell_rich)
                            else:
                                safe = cp_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                cell_parts.append(safe)
                        
                        cell_html = '<br/>'.join(cell_parts) if cell_parts else ''
                        
                        cell_style = ParagraphStyle(
                            name=f'cell_{id(cell)}',
                            fontName=_font_base,
                            fontSize=10,
                            leading=13,
                            alignment=TA_LEFT,
                        )
                        
                        try:
                            row_data.append(Paragraph(cell_html, cell_style))
                        except Exception:
                            row_data.append(Paragraph(cell.text.strip()[:100], cell_style))
                    
                    table_data.append(row_data)
                
                if table_data:
                    col_count = max(len(row) for row in table_data)
                    # Qatorlarni bir xil uzunlikka keltirish
                    for row in table_data:
                        while len(row) < col_count:
                            row.append(Paragraph('', ParagraphStyle(name=f'empty_{id(row)}', fontName=_font_base, fontSize=10)))
                    
                    col_widths = [available_width / col_count] * col_count
                    
                    pdf_table = Table(table_data, colWidths=col_widths)
                    pdf_table.setStyle(TableStyle([
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                        ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.92, 0.92, 0.96)),
                        ('FONTNAME', (0, 0), (-1, 0), _font_bold),
                        ('FONTSIZE', (0, 0), (-1, -1), 10),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('TOPPADDING', (0, 0), (-1, -1), 4),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                        ('LEFTPADDING', (0, 0), (-1, -1), 6),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ]))
                    story.append(Spacer(1, 6))
                    story.append(pdf_table)
                    story.append(Spacer(1, 6))
            except Exception as e:
                print(f"Jadval qayta ishlashda xatolik: {e}")
    
    if not story:
        story.append(Paragraph("(Bo'sh hujjat)", ParagraphStyle(
            name='empty_doc', fontName=_font_base, fontSize=12, alignment=TA_CENTER)))
    
    try:
        pdf_doc.build(story)
    except Exception as e:
        print(f"PDF build xatolik: {e}")
        import traceback
        traceback.print_exc()
        # Oddiy fallback — faqat matn
        story_simple = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                safe = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story_simple.append(Paragraph(safe, ParagraphStyle(
                    name=f'simple_{id(para)}', fontName=_font_base, fontSize=12, leading=16)))
            else:
                story_simple.append(Spacer(1, 6))
        if not story_simple:
            story_simple.append(Paragraph("(Bo'sh hujjat)", ParagraphStyle(
                name='simple_empty', fontName=_font_base, fontSize=12, alignment=TA_CENTER)))
        pdf_doc2 = SimpleDocTemplate(pdf_path, pagesize=A4,
            leftMargin=25*mm, rightMargin=25*mm, topMargin=20*mm, bottomMargin=20*mm)
        pdf_doc2.build(story_simple)
    
    return pdf_path


def _add_qr_overlay(request, order, pdf_path, temp_files):
    """PDF ustiga QR kodlar qo'shish.
    
    Joylashuvi (screenshot bo'yicha):
    - Top-right: Hujjat QR kodi (download link)
    - Left margin: Har bir imzolovchi uchun alohida QR kod (vertikal tarzda)
    """
    import tempfile
    import os
    from PyPDF2 import PdfReader, PdfWriter
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    
    final_buffer = BytesIO()
    
    try:
        # QR-kod overlay yaratish
        overlay_buffer = BytesIO()
        c = canvas.Canvas(overlay_buffer, pagesize=A4)
        width, height = A4
        
        # === TOP RIGHT: Hujjat QR (Download Link) ===
        download_url = request.build_absolute_uri(f"/documents/download-pdf/{order.id}/")
        qr = qrcode.QRCode(version=1, box_size=5, border=1)
        qr.add_data(download_url)
        qr.make(fit=True)
        img = qr.make_image(fill='black', back_color='white')
        
        fd, main_qr_path = tempfile.mkstemp(suffix='.png')
        os.close(fd)
        temp_files.append(main_qr_path)
        img.save(main_qr_path, format='PNG')
        
        # Top-right joylash (Document umumiy QR)
        qr_size = 90.0
        c.drawImage(main_qr_path, float(width - qr_size - 20), float(height - qr_size - 20),
                     width=qr_size, height=qr_size)
        
        # === BOTTOM LEFT: Xodimlar va Tasdiq QR kodi ===
        signatures = order.signatures.filter(signed=True).order_by('order_number')
        
        left_x = 30.0  # Chap chekkadan masofa
        bottom_y = 30.0  # Eng pastdan masofa
        bottom_qr_size = 90.0
        
        # Determine the bottom QR code to draw (use director's final QR if available, else the main doc QR)
        bottom_qr_path = order.final_qr_code.path if (order.director_approved and order.final_qr_code) else main_qr_path
        
        try:
            c.drawImage(bottom_qr_path, left_x, bottom_y, width=bottom_qr_size, height=bottom_qr_size)
        except Exception as e:
            print(f"Error drawing bottom QR: {e}")
            
        # Draw employee info next to the bottom QR
        text_x = left_x + bottom_qr_size + 15.0
        current_text_y = bottom_y + bottom_qr_size - 10.0
        c.setFont("Helvetica", 8)
        
        def safe_str(s):
            return str(s or '').encode('latin-1', 'replace').decode('latin-1')
            
        if order.director_approved and order.director_approved_at:
            c.drawString(text_x, float(current_text_y), f"Direktor tasdiqladi: {order.director_approved_at.strftime('%d.%m.%Y %H:%M')}")
            current_text_y -= 15.0
            
        if signatures.exists():
            c.setFont("Helvetica-Bold", 8)
            c.drawString(text_x, float(current_text_y), "Elektron imzolar (Xodimlar):")
            current_text_y -= 12.0
            c.setFont("Helvetica", 8)
            
            for i, sig in enumerate(signatures, 1):
                user = sig.user
                first_name = user.first_name
                last_name = user.last_name
                middle_name = user.middle_name if hasattr(user, 'middle_name') else ''
                position = user.position or '-'
                
                c.drawString(text_x, float(current_text_y), f"{i}. {safe_str(last_name)} {safe_str(first_name)} {safe_str(middle_name)}         {safe_str(position)}         {sig.signed_at.strftime('%d.%m.%Y %H:%M') if sig.signed_at else '-'}")
                # current_text_y -= 10.0
                # c.drawString(text_x + 10.0, float(current_text_y), f"Lavozim: {safe_str(position)}")
                # current_text_y -= 10.0
                # c.drawString(text_x + 10.0, float(current_text_y), f"Sana: {sig.signed_at.strftime('%d.%m.%Y %H:%M') if sig.signed_at else '-'}")
                current_text_y -= 15.0
        
        c.save()
        overlay_buffer.seek(0)
        
        # Asosiy PDF bilan overlay'ni birlashtirish
        overlay_pdf = PdfReader(overlay_buffer)
        overlay_page = overlay_pdf.pages[0]
        
        pdf_writer = PdfWriter()
        original_reader = PdfReader(pdf_path)
        
        for page_num in range(len(original_reader.pages)):
            page = original_reader.pages[page_num]
            if page_num == 0:
                # QR kodlarni faqat birinchi sahifaga qo'yish
                page.merge_page(overlay_page)
            pdf_writer.add_page(page)
        
        pdf_writer.write(final_buffer)
        
    except Exception as e:
        print(f"QR overlay error: {e}")
        import traceback
        traceback.print_exc()
        # Xato bo'lsa, original PDF ni qaytarish
        final_buffer = BytesIO()
        with open(pdf_path, 'rb') as f:
            final_buffer.write(f.read())
    
    final_buffer.seek(0)
    return final_buffer

@login_required
def director_approve_page(request, order_id):
    order = get_object_or_404(Order, id=order_id)
    if request.user.role != 'director':
        messages.error(request, "Faqat direktor tasdiqlashi mumkin.")
        return redirect('dashboard')
    
    if not order.is_fully_signed():
        messages.error(request, "Barcha xodimlar imzolab bo'lmagan.")
        return redirect('order_detail', order_id=order.id)
        
    if order.director_approved:
        messages.info(request, "Bu hujjat allaqachon tasdiqlangan.")
        return redirect('order_detail', order_id=order.id)
        
    return render(request, 'documents/director_approve.html', {'order': order})

def stamp_pdf_with_qrs(original_file, employee_qr_path, director_qr_paths=None, all_employee_qrs=None, employee_info_list=None):
    """
    Helper to stamp a PDF with an employee QR code(s) at the bottom left, 
    and optionally director QR codes at the top left.
    Returns a BytesIO buffer with the stamped PDF.
    """
    import tempfile, os
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas as rl_canvas
    from PyPDF2 import PdfReader, PdfWriter
    from io import BytesIO

    fd, temp_pdf = tempfile.mkstemp(suffix='.pdf')
    os.close(fd)
    
    with open(temp_pdf, 'wb') as f:
        f.write(original_file.read())
        
    original_file.seek(0)
    
    overlay_buffer = BytesIO()
    c = rl_canvas.Canvas(overlay_buffer, pagesize=A4)
    width, height = A4
    
    # Increase QR Size
    qr_size = 90.0
    margin = 30.0
    
    # Employee QR codes (bottom left)
    # the user wants ONE QR code at the bottom left with text stacked next to it.
    employee_data = employee_info_list if employee_info_list else []
    
    # We will draw the FIRST employee's QR code (or the only one) at the bottom left.
    bottom_qr_path = employee_qr_path
    if not bottom_qr_path and employee_data and employee_data[0].get('qr_path'):
        bottom_qr_path = employee_data[0]['qr_path']
        
    x_pos_emp = float(margin)
    y_pos_emp = float(margin)
    
    if bottom_qr_path and os.path.exists(bottom_qr_path):
        c.drawImage(bottom_qr_path, x_pos_emp, y_pos_emp, width=qr_size, height=qr_size)
    
    # Draw text next to the BOTTOM QR
    text_x = x_pos_emp + qr_size + 15.0
    current_text_y = y_pos_emp + qr_size - 10.0
    
    def safe_str(s):
        return str(s or '').encode('latin-1', 'replace').decode('latin-1')
        
    if employee_data:
        c.setFont("Helvetica-Bold", 8)
        c.drawString(text_x, float(current_text_y), "Elektron imzolar (Xodimlar):")
        current_text_y -= 12.0
        c.setFont("Helvetica", 8)
        
        for i, emp in enumerate(employee_data, 1):
            c.drawString(text_x, float(current_text_y), f"{i}. F.I.O: {safe_str(emp.get('full_name'))}")
            current_text_y -= 10.0
            c.drawString(text_x + 10.0, float(current_text_y), f"Lavozim: {safe_str(emp.get('position'))}")
            current_text_y -= 10.0
            c.drawString(text_x + 10.0, float(current_text_y), f"Sana: {safe_str(emp.get('date'))}")
            current_text_y -= 15.0
        
    # Director / Main QR code (top left)
    if director_qr_paths and len(director_qr_paths) > 0 and os.path.exists(director_qr_paths[0]):
        x_pos_dir = float(margin)
        y_pos_dir = float(height) - float(qr_size) - float(margin)
        c.drawImage(director_qr_paths[0], x_pos_dir, y_pos_dir, width=qr_size, height=qr_size)

    c.save()
    overlay_buffer.seek(0)
    
    overlay_reader = PdfReader(overlay_buffer)
    overlay_page = overlay_reader.pages[0]
    
    pdf_writer = PdfWriter()
    original_reader = PdfReader(temp_pdf)
    
    for page_num in range(len(original_reader.pages)):
        page = original_reader.pages[page_num]
        if page_num == 0:
            page.merge_page(overlay_page)
        pdf_writer.add_page(page)
        
    final_buffer = BytesIO()
    pdf_writer.write(final_buffer)
    final_buffer.seek(0)
    
    if os.path.exists(temp_pdf):
        try:
            os.remove(temp_pdf)
        except Exception:
            pass
            
    return final_buffer

def stamp_word_with_qrs(original_file, employee_qr_path, director_qr_paths=None, all_employee_qrs=None, employee_info_list=None):
    """
    Helper to stamp a Word document with QR codes.
    Appends the employee QR codes at the end.
    If director QR is given, adds it at the beginning.
    Returns a BytesIO buffer.
    """
    import tempfile, os
    import docx
    from docx.shared import Inches
    from io import BytesIO

    fd, temp_word = tempfile.mkstemp(suffix='.docx')
    os.close(fd)
    
    with open(temp_word, 'wb') as f:
        f.write(original_file.read())
        
    original_file.seek(0)
    
    doc_obj = docx.Document(temp_word)
    
    # Prepend director QR
    if director_qr_paths and len(director_qr_paths) > 0 and os.path.exists(director_qr_paths[0]):
        p_top = doc_obj.paragraphs[0].insert_paragraph_before()
        r_top = p_top.add_run()
        r_top.add_picture(director_qr_paths[0], width=Inches(1.5))
        p_top.add_run(" Elektron Nusxa (Umumiy Tasdiq)")
    
    # Append employee info next to one QR
    employee_data = employee_info_list if employee_info_list else []
    
    bottom_qr_path = employee_qr_path
    if not bottom_qr_path and employee_data and employee_data[0].get('qr_path'):
        bottom_qr_path = employee_data[0]['qr_path']
        
    if employee_data or bottom_qr_path:
        doc_obj.add_paragraph("--- Elektron Imzolar (Xodimlar) ---")
        
        # We can implement 'next to' using a simple table in Word
        table = doc_obj.add_table(rows=1, cols=2)
        table.autofit = True
        
        cell_img = table.cell(0, 0)
        cell_text = table.cell(0, 1)
        
        if bottom_qr_path and os.path.exists(bottom_qr_path):
            p_img = cell_img.paragraphs[0]
            r_img = p_img.add_run()
            r_img.add_picture(bottom_qr_path, width=Inches(1.5))
            
        p_text = cell_text.paragraphs[0]
        for i, emp in enumerate(employee_data, 1):
            if emp.get('full_name'):
                p_text.add_run(f"{i}. F.I.O: {emp.get('full_name')}\n").bold = True
            if emp.get('position'):
                p_text.add_run(f"    Lavozim: {emp.get('position')}\n")
            if emp.get('date'):
                p_text.add_run(f"    Sana: {emp.get('date')}\n\n")
        
    final_buffer = BytesIO()
    doc_obj.save(final_buffer)
    final_buffer.seek(0)
    
    if os.path.exists(temp_word):
        try:
            os.remove(temp_word)
        except Exception:
            pass
            
    return final_buffer

@require_http_methods(["POST"])
@login_required
def api_director_approve(request, order_id):
    order = get_object_or_404(Order, id=order_id)
    if request.user.role != 'director':
        return JsonResponse({'error': "Ruxsat yo'q"}, status=403)
        
    if not order.is_fully_signed():
        return JsonResponse({'error': "Barcha xodimlar imzolamagan"}, status=400)
        
    if order.director_approved:
        return JsonResponse({'error': "Allaqachon tasdiqlangan"}, status=400)
        
    # Tasdiqlash
    order.director_approved = True
    order.director_approved_at = timezone.now()
    
    # Umumiy QR kod yaratish
    signatures = order.signatures.all().order_by('order_number')
    qr_data_lines = [f"Buyruq №: {order.number}"]
    qr_data_lines.append(f"Direktor: {request.user.get_full_name()} (Tasdiq: {order.director_approved_at.strftime('%d.%m.%Y %H:%M')})")
    qr_data_lines.append("--- Imzo chekkanlar ---")
    
    for sig in signatures:
        time_str = sig.signed_at.strftime('%d.%m.%Y %H:%M') if sig.signed_at else '-'
        pos = sig.user.position or 'Xodim'
        qr_data_lines.append(f"{sig.user.get_full_name()} ({pos}) - {time_str}")
        
    qr_data = "\n".join(qr_data_lines)
    
    qr = qrcode.QRCode(version=1, box_size=10, border=5)
    qr.add_data(qr_data)
    qr.make(fit=True)
    img = qr.make_image(fill='black', back_color='white')
    
    buffer = BytesIO()
    img.save(buffer, format='PNG')
    order.final_qr_code.save(f"final_qr_{order.id}.png", File(buffer), save=False)
    order.save()
    
    # Asosiy buyruq uchun PDF generatsiya va QR larni yopishtirish
    # Odatda bu download paytida yaratiladi, lekin avtomatik qilish ham mumkin.
    # Bu yassi yerda hozircha asosiy order_pdf logic'ini o'zgarishsiz qoldiramiz 
    # MCHJ uchun order.stamped_file field bizda yo'q, shuning uchun order.pdf yaratish upload_stamped_pdf da yoki download da ishlaydi.
    
    # === Qo'shimcha Hujjatlar uchun Pechatli PDF Yaratish ===
    import os
    import tempfile
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas as rl_canvas
    from PyPDF2 import PdfReader, PdfWriter
    
    additional_docs = order.additional_docs.filter(is_signed=True).exclude(file='')
    for doc in additional_docs:
        # doc.file, doc.qr_code, order.final_qr_code bor
        if not doc.qr_code or not order.final_qr_code:
            continue
            
        temp_files_doc = []
        try:
            is_pdf = doc.file.name.lower().endswith('.pdf')
            is_word = doc.file.name.lower().endswith(('.doc', '.docx'))

            if not is_pdf and not is_word:
                continue
            
            # Qo'shimcha hujjat uchun yuklab olish URL-i asosida alohida QR yasash
            download_url = request.build_absolute_uri(f"/documents/download-additional-doc/{doc.id}/")
            qr_add = qrcode.QRCode(version=1, box_size=5, border=1)
            qr_add.add_data(download_url)
            qr_add.make(fit=True)
            img_add = qr_add.make_image(fill='black', back_color='white')
            
            fd, main_add_qr_path = tempfile.mkstemp(suffix='.png')
            os.close(fd)
            temp_files_doc.append(main_add_qr_path)
            img_add.save(main_add_qr_path, format='PNG')
                
            final_buffer = None
            
            employee_info = []
            if doc.qr_code and doc.signer:
                signer_user = doc.signer
                full_name = f"{signer_user.last_name} {signer_user.first_name}"
                if hasattr(signer_user, 'middle_name') and signer_user.middle_name:
                    full_name += f" {signer_user.middle_name}"
                employee_info.append({
                    'qr_path': doc.qr_code.path,
                    'full_name': full_name,
                    'position': signer_user.position or '-',
                    'date': doc.signed_at.strftime('%d.%m.%Y %H:%M') if doc.signed_at else '-'
                })
                
            if is_pdf:
                # `doc.qr_code` is the employee's signature QR code, `main_add_qr_path` is the director/system QR for the document
                final_buffer = stamp_pdf_with_qrs(
                    original_file=doc.file, 
                    employee_qr_path=doc.qr_code.path if doc.qr_code else None, 
                    director_qr_paths=[main_add_qr_path],
                    employee_info_list=employee_info
                )
                stamped_filename = f"stamped_doc_{doc.id}_{order.number}.pdf"
            elif is_word:
                final_buffer = stamp_word_with_qrs(
                    original_file=doc.file, 
                    employee_qr_path=doc.qr_code.path if doc.qr_code else None, 
                    director_qr_paths=[main_add_qr_path],
                    employee_info_list=employee_info
                )
                stamped_filename = f"stamped_doc_{doc.id}_{order.number}.docx"
                
            if final_buffer:
                doc.stamped_file.save(stamped_filename, File(final_buffer), save=False)
                doc.save()
            
        except Exception as e:
            print(f"Error generating stamped additional doc {doc.id}: {e}")
        finally:
            for tf in temp_files_doc:
                if os.path.exists(tf):
                    try:
                        os.remove(tf)
                    except:
                        pass
    
    Notification.objects.create(
        user=order.created_by,
        notification_type='status_changed',
        order=order,
        message=f"Sizning '{order.title}' buyrug'ingiz direktor tomonidan tasdiqlandi"
    )
    
    return JsonResponse({
        'success': True,
        'qr_code_url': order.final_qr_code.url if order.final_qr_code else ''
    })

@login_required
def delete_order(request, order_id):
    if request.user.role != 'admin':
        messages.error(request, 'Ruxsat yo\'q')
        return redirect('dashboard')
    
    order = get_object_or_404(Order, id=order_id)
    if request.method == 'POST':
        order.delete()
        messages.success(request, 'Hujjat muvaffaqiyatli o\'chirildi')
        return redirect('dashboard')
        
    return render(request, 'documents/confirm_delete.html', {'order': order})

